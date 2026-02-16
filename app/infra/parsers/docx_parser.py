"""
DOCX 파서 구현
"""
from typing import List
from .base import BaseParser
from .models import ParsedContent, Table, Image


class DOCXParser(BaseParser):
    """DOCX 파일 전용 파서"""
    
    def get_supported_mime_types(self) -> List[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
    
    async def parse(self, file_path: str) -> ParsedContent:
        """DOCX 파일에서 구조화된 콘텐츠를 추출합니다."""
        if not self.validate_file(file_path):
            raise ValueError(f"유효하지 않은 DOCX 파일: {file_path}")
        
        try:
            from docx import Document
            from docx.document import Document as DocumentType
            
            doc = Document(file_path)
            
            # 텍스트 추출 (구조 정보 포함)
            full_text = ""
            structure = {"headings": [], "paragraphs": [], "sections": []}
            tables = []
            images = []
            
            # 문단별 텍스트 및 구조 추출
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # 스타일 정보 확인
                style = para.style.name if para.style else "Normal"
                
                # 제목 구조 파악
                if style.startswith("Heading"):
                    level = self._extract_heading_level(style)
                    structure["headings"].append({
                        "level": level,
                        "text": text,
                        "style": style
                    })
                    full_text += f"{'#' * level} {text}\n\n"
                else:
                    structure["paragraphs"].append({
                        "text": text,
                        "style": style
                    })
                    full_text += f"{text}\n\n"
            
            # 테이블 추출
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                headers = []
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    
                    if row_idx == 0:  # 첫 번째 행을 헤더로 간주
                        headers = row_data
                    else:
                        table_data.append(row_data)
                
                if headers or table_data:
                    tables.append(Table(
                        headers=headers,
                        rows=table_data,
                        position={"table_index": table_idx},
                        metadata={
                            "extraction_method": "python-docx",
                            "row_count": len(table.rows),
                            "col_count": len(table.rows[0].cells) if table.rows else 0
                        }
                    ))
            
            # 이미지 정보 추출 (간단한 버전)
            # DOCX의 이미지는 복잡하므로 기본 정보만
            try:
                from docx.oxml.ns import qn
                
                # 문서의 part에서 이미지 관계 찾기
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        images.append(Image(
                            name=rel.target_ref.split("/")[-1],
                            format="unknown",  # DOCX에서 정확한 format 추출은 복잡
                            metadata={
                                "extraction_method": "python-docx",
                                "relationship_id": rel.rId
                            }
                        ))
            except Exception as e:
                self.logger.warning(f"DOCX 이미지 정보 추출 실패: {e}")
            
            # 문서 메타데이터 추출
            metadata = {
                "parser": "python-docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "heading_count": len(structure["headings"])
            }
            
            # 핵심 속성 추출
            try:
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else ""
                })
            except Exception as e:
                self.logger.warning(f"DOCX 메타데이터 추출 실패: {e}")
            
            return ParsedContent(
                raw_text=full_text.strip(),
                metadata=metadata,
                structure=structure,
                tables=tables,
                images=images
            )
            
        except ImportError:
            raise ImportError("python-docx가 설치되지 않았습니다: pip install python-docx")
        except Exception as e:
            raise ValueError(f"DOCX 파싱 실패: {e}")
    
    def _extract_heading_level(self, style_name: str) -> int:
        """제목 스타일에서 레벨 추출"""
        if "Heading" in style_name:
            try:
                # "Heading 1" -> 1, "Heading 2" -> 2 등
                level_str = style_name.split()[-1]
                return int(level_str)
            except (ValueError, IndexError):
                return 1
        return 1