import hashlib
import logging
from typing import List, Dict, Any

from langchain_core.documents import Document as LangChainDocument
# LangChain text splitters
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.core.config import settings


class DocumentProcessor:
    """문서 처리 및 청킹 서비스"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

        # 환경변수 기반 텍스트 분할기 설정
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=settings.CHUNK_SEPARATORS_LIST
        )

    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """
        파일에서 텍스트를 추출합니다.
        
        Args:
            file_path: 파일 경로
            file_type: 파일 MIME 타입
            
        Returns:
            str: 추출된 텍스트
        """
        try:
            if file_type == "text/plain":
                return self._extract_from_txt(file_path)
            elif file_type == "application/pdf":
                return self._extract_from_pdf(file_path)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"지원하지 않는 파일 형식: {file_type}")

        except Exception as e:
            self.logger.error(f"텍스트 추출 실패 ({file_path}): {e}")
            raise

    def _extract_from_txt(self, file_path: str) -> str:
        """TXT 파일에서 텍스트 추출"""
        encodings = ['utf-8', 'cp949', 'latin-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        raise ValueError(f"텍스트 파일 인코딩을 식별할 수 없습니다: {file_path}")

    def _extract_from_pdf(self, file_path: str) -> str:
        """PDF 파일에서 텍스트 추출 (LangChain PDFLoader 사용)"""
        try:
            from langchain_community.document_loaders import PyMuPDFLoader

            # PyMuPDFLoader 사용 (가장 안정적)
            loader = PyMuPDFLoader(file_path)
            documents = loader.load()

            # 모든 페이지의 텍스트 결합
            text = ""
            for doc in documents:
                page_text = doc.page_content.strip()
                if page_text:
                    text += page_text + "\n\n"

            return text.strip()

        except ImportError:
            # PyMuPDF가 없으면 PyPDF2 fallback
            self.logger.warning("PyMuPDF가 없습니다. PyPDF2로 fallback합니다.")
            return self._extract_from_pdf_fallback(file_path)
        except Exception as e:
            self.logger.warning(f"PyMuPDF 추출 실패, PyPDF2로 fallback: {e}")
            return self._extract_from_pdf_fallback(file_path)

    def _extract_from_pdf_fallback(self, file_path: str) -> str:
        """PDF 파일에서 텍스트 추출 (PyPDF2 fallback)"""
        try:
            import PyPDF2

            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text += page.extract_text() + "\n"
                    except Exception as e:
                        self.logger.warning(f"PDF 페이지 {page_num} 추출 실패: {e}")
                        continue

            return text.strip()

        except ImportError:
            raise ImportError("PDF 처리를 위해 pymupdf 또는 PyPDF2가 필요합니다: pip install pymupdf 또는 pip install PyPDF2")
        except Exception as e:
            raise ValueError(f"PDF 텍스트 추출 실패: {e}")

    def _extract_from_docx(self, file_path: str) -> str:
        """DOCX 파일에서 텍스트 추출"""
        try:
            from docx import Document

            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return text.strip()

        except ImportError:
            raise ImportError("DOCX 처리를 위해 python-docx가 필요합니다: pip install python-docx")
        except Exception as e:
            raise ValueError(f"DOCX 텍스트 추출 실패: {e}")

    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        텍스트를 청크로 분할합니다.
        
        Args:
            text: 분할할 텍스트
            metadata: 추가 메타데이터
            
        Returns:
            List[Dict]: 청크 정보 리스트
        """
        if not text.strip():
            return []

        try:
            # LangChain Document 생성
            doc = LangChainDocument(
                page_content=text,
                metadata=metadata or {}
            )

            # 텍스트 분할
            chunks = self.text_splitter.split_documents([doc])

            # 청크 정보 생성
            chunk_data = []
            for i, chunk in enumerate(chunks):
                content = chunk.page_content.strip()
                if not content:
                    continue

                chunk_info = {
                    "chunk_index": i,
                    "content": content,
                    "content_hash": self._generate_content_hash(content),
                    "char_count": len(content),
                    "token_count": self._estimate_token_count(content),
                    "chunk_type": "text",
                    "start_position": text.find(content) if content in text else 0,
                    "end_position": text.find(content) + len(content) if content in text else len(content),
                    "extra_metadata": chunk.metadata
                }

                chunk_data.append(chunk_info)

            self.logger.info(f"텍스트를 {len(chunk_data)}개 청크로 분할 완료")
            return chunk_data

        except Exception as e:
            self.logger.error(f"텍스트 청킹 실패: {e}")
            raise

    def _generate_content_hash(self, content: str) -> str:
        """콘텐츠 해시 생성"""
        return hashlib.sha256(content.encode()).hexdigest()

    def _estimate_token_count(self, text: str) -> int:
        """토큰 수 추정 (대략적)"""
        # 간단한 추정: 단어 수 * 1.3 (한국어 고려)
        words = len(text.split())
        korean_chars = len([c for c in text if ord(c) >= 0xAC00 and ord(c) <= 0xD7A3])

        # 한국어 문자는 토큰 비율이 높음
        estimated_tokens = words * 1.3 + korean_chars * 0.5
        return int(estimated_tokens)

    def validate_file_content(self, text: str) -> bool:
        """
        추출된 텍스트의 유효성을 검증합니다.
        
        Args:
            text: 검증할 텍스트
            
        Returns:
            bool: 유효성 여부
        """
        if not text or not text.strip():
            return False

        # 최소 길이 체크
        if len(text.strip()) < 10:
            return False

        # 의미있는 문자 비율 체크
        meaningful_chars = sum(1 for c in text if c.isalnum() or c in '.,!?;:')
        if len(text) > 0 and meaningful_chars / len(text) < 0.1:
            return False

        return True
