import csv
import io
import json
import mimetypes
import time
from datetime import datetime
from pathlib import Path

import PyPDF2
import pandas as pd
from docx import Document

from .schemas import (
    FileType, BaseFileMetadata, FileExtractionResult,
    PDFExtractionResult, PDFMetadata,
    DocxExtractionResult, DocxMetadata,
    ExcelExtractionResult, ExcelMetadata, ExcelSheetInfo,
    CSVExtractionResult, CSVMetadata,
    TextExtractionResult, TextMetadata,
    JSONExtractionResult, JSONMetadata,
    ImageExtractionResult, ImageMetadata
)


class FileExtractor:
    """순수한 파일 내용 추출 유틸리티"""

    SUPPORTED_TYPES = {
        'application/pdf': FileType.PDF,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': FileType.DOCX,
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': FileType.XLSX,
        'application/vnd.ms-excel': FileType.XLS,
        'text/plain': FileType.TXT,
        'text/csv': FileType.CSV,
        'application/json': FileType.JSON,
        'image/jpeg': FileType.IMAGE,
        'image/png': FileType.IMAGE,
        'image/gif': FileType.IMAGE,
    }

    @staticmethod
    def detect_file_type(file_content: bytes, filename: str = None) -> str:
        """파일 타입 감지"""
        # 파일명 확장자로 MIME 타입 감지
        if filename:
            guessed_type, _ = mimetypes.guess_type(filename)
            if guessed_type:
                return guessed_type

            # 확장자 기반 추가 매핑
            ext = Path(filename).suffix.lower()
            ext_mapping = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.xls': 'application/vnd.ms-excel',
                '.csv': 'text/csv',
                '.txt': 'text/plain',
                '.json': 'application/json',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.gif': 'image/gif',
            }
            if ext in ext_mapping:
                return ext_mapping[ext]

        # 파일 내용 기반 간단한 감지
        if file_content.startswith(b'%PDF'):
            return 'application/pdf'
        elif file_content.startswith(b'PK'):  # ZIP 기반 파일 (docx, xlsx)
            if filename and filename.endswith('.docx'):
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif filename and filename.endswith('.xlsx'):
                return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif file_content.startswith((b'\xff\xd8\xff', b'\x89PNG')):  # JPEG, PNG
            return 'image/jpeg' if file_content.startswith(b'\xff\xd8\xff') else 'image/png'

        # 기본값
        return 'application/octet-stream'

    @staticmethod
    def extract_pdf_content(file_content: bytes) -> PDFExtractionResult:
        """PDF 파일 내용 추출"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_content = ""
            page_count = len(pdf_reader.pages)

            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"

            metadata_raw = pdf_reader.metadata if pdf_reader.metadata else {}

            pdf_metadata = PDFMetadata(
                title=metadata_raw.get("/Title", ""),
                author=metadata_raw.get("/Author", ""),
                subject=metadata_raw.get("/Subject", ""),
                creator=metadata_raw.get("/Creator", "")
            )

            return PDFExtractionResult(
                content=text_content.strip(),
                page_count=page_count,
                metadata=pdf_metadata,
                file_info=BaseFileMetadata(
                    mime_type="application/pdf",
                    file_type=FileType.PDF
                )
            )
        except Exception as e:
            raise ValueError(f"PDF 추출 실패: {str(e)}")

    @staticmethod
    def extract_docx_content(file_content: bytes) -> DocxExtractionResult:
        """DOCX 파일 내용 추출"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            # 단락별 텍스트 추출
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n".join(paragraphs)

            # 테이블 내용 추출
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append(table_data)

            docx_metadata = DocxMetadata(
                paragraph_count=len(paragraphs),
                table_count=len(tables_content)
            )

            return DocxExtractionResult(
                content=full_text,
                paragraphs=paragraphs,
                tables=tables_content,
                metadata=docx_metadata,
                file_info=BaseFileMetadata(
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    file_type=FileType.DOCX
                )
            )
        except Exception as e:
            raise ValueError(f"DOCX 추출 실패: {str(e)}")

    @staticmethod
    def extract_excel_content(file_content: bytes, sheet_name: str = None) -> ExcelExtractionResult:
        """Excel 파일 내용 추출"""
        try:
            excel_file = io.BytesIO(file_content)

            # 모든 시트 이름 가져오기
            xl_file = pd.ExcelFile(excel_file)
            sheet_names = xl_file.sheet_names

            sheets_data = {}

            if sheet_name and sheet_name in sheet_names:
                # 특정 시트만 읽기
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                sheets_data[sheet_name] = ExcelSheetInfo(
                    data=df.to_dict('records'),
                    columns=df.columns.tolist(),
                    shape=df.shape
                )
            else:
                # 모든 시트 읽기
                for sheet in sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet)
                    sheets_data[sheet] = ExcelSheetInfo(
                        data=df.to_dict('records'),
                        columns=df.columns.tolist(),
                        shape=df.shape
                    )

            excel_metadata = ExcelMetadata(total_sheets=len(sheet_names))

            return ExcelExtractionResult(
                content=sheets_data,
                sheet_names=sheet_names,
                metadata=excel_metadata,
                file_info=BaseFileMetadata(
                    mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    file_type=FileType.XLSX
                )
            )
        except Exception as e:
            raise ValueError(f"Excel 추출 실패: {str(e)}")

    @staticmethod
    def extract_csv_content(file_content: bytes, encoding: str = 'utf-8') -> CSVExtractionResult:
        """CSV 파일 내용 추출"""
        try:
            # 여러 인코딩 시도
            encodings = [encoding, 'utf-8', 'cp949', 'euc-kr', 'latin1']

            for enc in encodings:
                try:
                    csv_text = file_content.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("지원되지 않는 인코딩입니다")

            # CSV 파싱
            csv_file = io.StringIO(csv_text)
            reader = csv.DictReader(csv_file)

            data = list(reader)
            columns = reader.fieldnames or []

            csv_metadata = CSVMetadata(
                row_count=len(data),
                column_count=len(columns),
                encoding_used=enc
            )

            return CSVExtractionResult(
                content=data,
                columns=columns,
                metadata=csv_metadata,
                file_info=BaseFileMetadata(
                    mime_type="text/csv",
                    file_type=FileType.CSV
                )
            )
        except Exception as e:
            raise ValueError(f"CSV 추출 실패: {str(e)}")

    @staticmethod
    def extract_text_content(file_content: bytes, encoding: str = 'utf-8') -> TextExtractionResult:
        """일반 텍스트 파일 내용 추출"""
        try:
            encodings = [encoding, 'utf-8', 'cp949', 'euc-kr', 'latin1']

            for enc in encodings:
                try:
                    text_content = file_content.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("지원되지 않는 인코딩입니다")

            lines = text_content.split('\n')

            text_metadata = TextMetadata(
                line_count=len(lines),
                character_count=len(text_content),
                encoding_used=enc
            )

            return TextExtractionResult(
                content=text_content,
                lines=lines,
                metadata=text_metadata,
                file_info=BaseFileMetadata(
                    mime_type="text/plain",
                    file_type=FileType.TXT
                )
            )
        except Exception as e:
            raise ValueError(f"텍스트 추출 실패: {str(e)}")

    @staticmethod
    def extract_json_content(file_content: bytes, encoding: str = 'utf-8') -> JSONExtractionResult:
        """JSON 파일 내용 추출"""
        try:
            json_text = file_content.decode(encoding)
            json_data = json.loads(json_text)

            json_metadata = JSONMetadata(
                data_type=type(json_data).__name__,
                size=len(json_text),
                encoding_used=encoding
            )

            return JSONExtractionResult(
                content=json_data,
                metadata=json_metadata,
                file_info=BaseFileMetadata(
                    mime_type="application/json",
                    file_type=FileType.JSON
                )
            )
        except Exception as e:
            raise ValueError(f"JSON 추출 실패: {str(e)}")

    @staticmethod
    def extract_content(file_content: bytes, filename: str = None, **kwargs) -> FileExtractionResult:
        """파일 내용 추출 메인 메서드"""
        start_time = time.time()

        mime_type = FileExtractor.detect_file_type(file_content, filename)
        file_type = FileExtractor.SUPPORTED_TYPES.get(mime_type)

        if not file_type:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {mime_type}")

        extractors = {
            FileType.PDF: FileExtractor.extract_pdf_content,
            FileType.DOCX: FileExtractor.extract_docx_content,
            FileType.XLSX: FileExtractor.extract_excel_content,
            FileType.XLS: FileExtractor.extract_excel_content,
            FileType.CSV: FileExtractor.extract_csv_content,
            FileType.TXT: FileExtractor.extract_text_content,
            FileType.JSON: FileExtractor.extract_json_content,
        }

        if file_type == FileType.IMAGE:
            image_metadata = ImageMetadata(size=len(file_content))
            return ImageExtractionResult(
                content=None,
                message="이미지 파일은 텍스트 추출을 지원하지 않습니다",
                metadata=image_metadata,
                file_info=BaseFileMetadata(
                    mime_type=mime_type,
                    file_type=FileType.IMAGE,
                    original_filename=filename,
                    file_size=len(file_content)
                )
            )

        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"추출기를 찾을 수 없습니다: {file_type}")

        result = extractor(file_content, **kwargs)

        # 공통 메타데이터 추가
        processing_time = time.time() - start_time
        result.file_info.original_filename = filename
        result.file_info.file_size = len(file_content)
        result.file_info.processing_time = processing_time
        result.file_info.processed_at = datetime.now()

        return result
